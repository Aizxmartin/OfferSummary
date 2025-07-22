
import streamlit as st
import pdfplumber
import tempfile
import os
from docx import Document
import re

def extract_fields_from_text(text):
    fields = {}

    def extract(pattern, label, flags=0, default="Not found"):
        match = re.search(pattern, text, flags)
        fields[label] = match.group(1).strip() if match else default

    # Buyer
    extract(r"Buyer\..*?\n(.*?)\s*\(Buyer\)", "2.1. Buyer Buyer(s) Names")

    # Title checkboxes logic
    if "☒ Other In Severalty" in text or "X Other In Severalty" in text:
        fields["2.1. Title Box Checked or In Severalty"] = "☒ Other – In Severalty"
    elif "☒ Joint Tenants" in text or "X Joint Tenants" in text:
        fields["2.1. Title Box Checked or In Severalty"] = "☒ Joint Tenants"
    elif "☒ Tenants In Common" in text or "X Tenants In Common" in text:
        fields["2.1. Title Box Checked or In Severalty"] = "☒ Tenants In Common"
    else:
        fields["2.1. Title Box Checked or In Severalty"] = "None Selected"

    # Inclusions and Exclusions
    extract(r"2\.5\.3.*?included in the Purchase Price:\s*(.*?)\n", "2.5.3. Other Inclusions")
    extract(r"2\.6\. Exclusions:\s*(.*?)\n", "2.6. Exclusions")

    # Deadlines
    extract(r"Time of Day Deadline\s*(.*?)\n", "3.1. Time of Day Deadline")
    extract(r"Alternative Earnest Money Deadline\s*(.*?)\n", "3.1. Alternative Earnest Money Deadline")
    extract(r"New Loan Terms Deadline\s*(.*?)\n", "3.1. New Loan Terms Deadline")
    extract(r"New Loan Availability Deadline\s*(.*?)\n", "3.1. New Loan Availability Deadline")
    extract(r"Inspection Termination Deadline\s*(.*?)\n", "3.1. Inspection Termination Deadline")
    extract(r"Closing Date\s*(.*?)\n", "3.1. Closing Date")
    extract(r"Possession Date\s*(.*?)\n", "3.1. Possession Date")
    extract(r"Possession Time\s*(.*?)\n", "3.1. Possession Time")

    # Money
    extract(r"Purchase Price.*?\$([\d,\.]+)", "3.1. Purchase Price")
    extract(r"Earnest Money.*?\$([\d,\.]+)", "4.1. Earnest Money")
    extract(r"Cash at Closing.*?\$([\d,\.]+)", "4.1. Cash at Closing")

    # Other monetary and obligation terms
    extract(r"Seller will credit to Buyer \$(.*?)\s*\(", "4.2 Seller Concession")
    extract(r"Earnest Money.*?in the form of a (.*?)\,", "4.3 Earnest held by")
    extract(r"Buyer.*?represents.*?(Does|Does Not)", "4.4.3. Available Funds")

    # Loan limitation
    if "☒ Conventional" in text or "X Conventional" in text:
        fields["4.5.3. Loan Limitations"] = "Conventional"
    elif "☒ FHA" in text or "X FHA" in text:
        fields["4.5.3. Loan Limitations"] = "FHA"
    elif "☒ VA" in text or "X VA" in text:
        fields["4.5.3. Loan Limitations"] = "VA"
    else:
        fields["4.5.3. Loan Limitations"] = "Not marked"

    extract(r"Cost of the Appraisal.*?paid by\s*(Buyer|Seller)", "6.4. Cost of Appraisal")

    # Title Selection
    fields["8.1.1. Seller Selects"] = "Yes" if "8.1.1" in text and ("☒" in text or "X" in text) else "No"
    fields["8.1.2. Buyer Selects"] = "Yes" if "8.1.2" in text and ("☒" in text or "X" in text) else "No"

    # Other Docs (10.6.1.6)
    extract(r"10\.6\.1\.6.*?Other documents and information:(.*?)\n10\.6\.2", "10.6.1.6. Other Docs", re.DOTALL)

    # Title transfer
    extract(r"deliver the following good and sufficient deed to Buyer, at Closing:\s*(.*?)\s", "13. Transfer of Title")

    # Section 29 commission %
    extract(r"29\.1.*?(\d\.\d+%)", "Section 29 (29.1/29.2/29.3)")

    # Additional provisions
    extract(r"30.*?Additional Provisions.*?Seller agrees to (.*?)\.", "Section 30 Additional Provisions", re.DOTALL)

    return fields

def main():
    st.title("Advanced Contract Summary Table Generator")
    uploaded = st.file_uploader("Upload Colorado Contract PDF", type="pdf")
    if uploaded:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded.read())
            tmp_path = tmp.name

        with pdfplumber.open(tmp_path) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])

        fields = extract_fields_from_text(text)

        # DOCX export
        doc = Document()
        doc.add_heading("Contract Summary Table", 0)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Extracted Value"
        for field, value in fields.items():
            row = table.add_row().cells
            row[0].text = field
            row[1].text = value

        output_path = os.path.join(tempfile.gettempdir(), "Contract_Summary_Table.docx")
        doc.save(output_path)
        with open(output_path, "rb") as file:
            st.download_button("Download Summary DOCX", file, file_name="Contract_Summary_Table.docx")

if __name__ == "__main__":
    main()
